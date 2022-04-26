"""
    Taylor J. Brown
    07MAR22
    file_Parse.py
    Opens a file and returns the parsed contents in a sorted dictionary
"""

import os
import csv
import docx
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def main():
    # File path (will be changed when intergrated with flask app)
    file_P_CSV = "CSC_289_CAP.csv"
    file_P_DOCX = "CSC_289_CAP.docx"
    DATA_FORMATTED_CSV = []
    DATA_FORMATTED_DOCX = []
    CLASS_NAME = ""
    
    
    CLASS_NAME, DATA_FORMATTED_CSV = ParseCSV(file_P_CSV)
    # Prints each DATA_FORMATTED_CSV value in the CSV file  
    print(CLASS_NAME)     
    for item in DATA_FORMATTED_CSV:
        print(item , "\n")
        
    print()
    
    CLASS_NAME, DATA_FORMATTED_DOCX = ParseDOCX(file_P_DOCX)
    # Prints the objects in the list
    print(CLASS_NAME) 
    for item in DATA_FORMATTED_DOCX:
        print(item , "\n") 


def ParseCSV(file_P_CSV):
    # Strips the file extention and assigns it as the Class_name
    size = len(file_P_CSV)
    CLASS_NAME = file_P_CSV[size - 20:size - 9]
    """
    Get the characters from position 2 to position 5 (not included):

    b = "Hello, World!"
    print(b[2:5])
    """
    ###MODIFY THIS TO THE CORRESPONDING PATH OF THE TEMPLATE FILE DIRECTORY!!!
    ###OR THE CLASSNAME IS NOT CORRECT!!!
    
    
    # Parses the rows and makes the first row read the header and the subsequent rows the values for each header.
    with open(file_P_CSV, "r") as file:
        csvreader = csv.DictReader(file)
        DATA_FORMATTED_CSV = []
        for row in csvreader:
            DATA_FORMATTED_CSV.append(row)
            
    return CLASS_NAME, DATA_FORMATTED_CSV
        

def ParseDOCX(file_P_DOCX):
    # Passes the file through the docx method
    document = Document(file_P_DOCX)
    
    # Chops off the file extention and assigns it to the class name
    size = len(file_P_DOCX)
    CLASS_NAME = file_P_DOCX[size - 20:size - 9]
    """
    Get the characters from position 2 to position 5 (not included):

    b = "Hello, World!"
    print(b[2:5])
    """
    ###MODIFY THIS TO THE CORRESPONDING PATH OF THE TEMPLATE FILE DIRECTORY!!!
    ###OR THE CLASSNAME IS NOT CORRECT!!!
    
    # Rips the data from the table and appends each item to the list
    data = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                data.append(cell.text)

    # Assigns the row headers as keys for the following records
    data_keys = []
    data_keys.append(data[0])
    data_keys.append(data[1])
    data_keys.append(data[2])
    
    # Iterates through the list and appends the dictionary key value pairs to the DATA_FORMATTED list 
    DATA_FORMATTED_DOCX = []
    count = int((len(data) / 3) - 1)
    index = 3
    for _ in range(count):
        tmp = {}
        tmp[data_keys[0]] = data[index]
        tmp[data_keys[1]] = data[index + 1]
        tmp[data_keys[2]] = data[index + 2]
        index += 3
        DATA_FORMATTED_DOCX.append(tmp)

    return CLASS_NAME, DATA_FORMATTED_DOCX

def files(num):
    #File name Section
    if num == 1:
        file = "CSC_289_CAP.csv"
        return file
    elif num == 2:
        file = "CSC_289_CAP.docx"
        return file
    elif num == 4:
        file = "JSONFILE.json"
        return file

def CreateCSV(path,file, CLASS_NAME, DATA_FORMATTED):
    #Creating CSV
    
    #### FOR WEB PATH STUFF
    #dir_path = os.path.abspath("/home/Pythonese/mysite/Filers")
    #File Directory Below
    dir_path = "C:\\Users\\lord_\\289-Captstone-Project-and-Development(tmp)\\File_Parser\\Filers"
    #Check Directory if exists and it doesn't create it
    os.makedirs(dir_path, exist_ok=True)
    #Joining the path with the file name
    with open(os.path.join(dir_path,file),"w+") as f:
        keys=DATA_FORMATTED[0].keys()
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        for key in DATA_FORMATTED:
            writer.writerow(key)


        # for row in csvreader:
        #     DATA_FORMATTED_CSV.append(row)
    
def CreateDOCX(path,file, CLASS_NAME, DATA_FORMATTED):
    #Create DOCX
    
    ### WEB PATH STUFF
    # dir_path = os.path.abspath("/home/Pythonese/mysite/Filers")
    #File Directory Below
    dir_path = "C:\\Users\\lord_\\289-Captstone-Project-and-Development(tmp)\\File_Parser\\Filers"
    #Check Directory if exists and it doesn't create it
    os.makedirs(dir_path, exist_ok=True)
    
    fname = os.path.join(dir_path, file)
    
    document = Document()
    
    b = document.add_paragraph()
    b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = b.add_run(CLASS_NAME)
    fonts = title.font
    fonts.size = Pt(18)
    fonts.bold = True
    
    
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Assignment'
    hdr_cells[1].text = 'Due-Date'
    hdr_cells[2].text = 'Comment'
    for item in DATA_FORMATTED:
        #print(item)
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['Assignment'])
        row_cells[1].text = str(item['Due-Date'])
        row_cells[2].text = str(item['Comment'])
        
    document.save(fname)

def Create_JSON(path,file, CLASS_NAME, DATA_FORMATTED):
    with open(file,"w+") as f:
        json.dump(DATA_FORMATTED,f, indent=3)

if __name__ == ("__main__"):
    main()
