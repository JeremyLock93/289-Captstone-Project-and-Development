#Lord Jenar Adolph C Allan
#
#CSC-289-0B01
import os
import File_Parser
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def DisplayMenu():
    print("File Creation")
    print("______________")
    print("1)	Create .docx file")
    print("2)	Create .csv file")
    print("3)	Exit")
    print()
def main():
    ''' This menu function is formatted to prevent error '''
    loop = '1'
    while loop == '1':
        DisplayMenu()
        selection = input("Choose one of the menu options: ")
        if selection == '1':
            num1 = 1
            file = files(num1)
            Create(file)
        elif selection == '2':
            num1 = 2
            file = files(num1)
            Create(file)      
        elif selection == '3':
            loop = '0'
            print()
            print("Bye!")
            ##raise SystemExit(0)#exit feature for spyder#REMOVED
        else:
            print("Invalid input/choice. Choose within 1-4")
        print()

def files(num):
    if num == 1:
        file = "CSC_289_CAP.csv"
        return file
    elif num == 2:
        file = "CSC_289_CAP.docx"
        return file
    
def CreateCSV(path,file, CLASS_NAME, DATA_FORMATTED):
    fname = os.path.join(path, file)
    with open(fname,"w+") as f:
        pass
            
def CreateDOCX(path,file, CLASS_NAME, DATA_FORMATTED):
    document = Document()
    fname = os.path.join(path, file)
    
    
    b = document.add_paragraph()
    b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = b.add_run(CLASS_NAME)
    fonts = title.font
    fonts.size = Pt(18)
    fonts.bold = True
    
    
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Assignment'
    hdr_cells[1].text = 'Due-Date'
    hdr_cells[2].text = 'Comment'
    for item in DATA_FORMATTED:
        #print(item)
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['Assignment'])
        row_cells[1].text = str(item['Due-Date'])
        row_cells[2].text = str(item['Comment'])
        
    document.save(fname)
            
def Read(file):
    #read file for verification
    with open(file,"r") as f:       
        for line in f.readlines():
            print(line)
            """ for word in line.split():
               print(word)  """  

if __name__== "__main__":
  main()